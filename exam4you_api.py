import streamlit as st
import time
import json
from PyPDF2 import PdfReader
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import io
import base64
from io import BytesIO
from openai import OpenAI

# Set Streamlit page configuration
st.set_page_config(page_title="Exam Creator", page_icon="📝", layout="wide")

__version__ = "1.5.0"  # Updated version number

# --------------------------- Helper Functions ---------------------------

def extract_text_from_pdf(pdf_file):
    """Extracts text content from an uploaded PDF file."""
    try:
        pdf_reader = PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(file):
    """Extracts text from a DOCX file."""
    try:
        doc = Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return ""

def process_image(file):
    """Processes an uploaded image and converts it to a Base64-encoded string."""
    try:
        image = Image.open(file)
        if image.mode != "RGB":
            image = image.convert("RGB")
        max_size = 1000
        if max(image.size) > max_size:
            image.thumbnail((max_size, max_size))
        buffered = io.BytesIO()
        image.save(buffered, format="JPEG")
        image_bytes = buffered.getvalue()
        return base64.b64encode(image_bytes).decode("utf-8")
    except Exception as e:
        st.error(f"Error processing image: {e}")
        return None

def chunk_text(text, max_tokens=3000):
    """Splits the extracted text into manageable chunks."""
    sentences = text.split('. ')
    chunks = []
    chunk = ""
    for sentence in sentences:
        if len(chunk) + len(sentence) > max_tokens:
            chunks.append(chunk)
            chunk = sentence + ". "
        else:
            chunk += sentence + ". "
    if chunk:
        chunks.append(chunk)
    return chunks

def generate_mc_questions(client, content_text, model):
    """Generates multiple-choice questions using OpenAI's API."""
    system_prompt = (
        "You are an educator tasked with creating a high school-level multiple-choice exam. "
        "Use the given content to generate single-choice questions. "
        "Each question must have one correct answer. Generate as many as necessary, up to 20 questions. "
        "Return the output as valid JSON with the structure: [{'question': '...', 'choices': ['...'], "
        "'correct_answer': '...', 'explanation': '...'}, ...]."
    )
    user_prompt = f"Content:\n\n{content_text}\n\nGenerate exam questions."

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]

    try:
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.5,
            max_tokens=16000
        )
        return response.choices[0].message.content, None
    except Exception as e:
        return None, f"Error generating questions: {e}"

def parse_generated_questions(response):
    """Parses the JSON response from OpenAI into Python objects."""
    try:
        json_start = response.find('[')
        json_end = response.rfind(']') + 1
        if json_start == -1 or json_end == 0:
            return None, f"No JSON data found in the response:\n{response[:500]}..."
        json_str = response[json_start:json_end]
        questions = json.loads(json_str)
        return questions, None
    except json.JSONDecodeError as e:
        return None, f"JSON parsing error: {e}\nResponse snippet:\n{response[:500]}..."
    except Exception as e:
        return None, f"Unexpected error: {str(e)}"

# --------------------------- PDF and DOCX Generation ---------------------------

def generate_pdf(questions, include_answers=True):
    """Generates a PDF file with exam questions."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, 'Generated Exam', 0, 1, 'C')

    pdf.set_font('Arial', '', 12)
    for i, q in enumerate(questions):
        pdf.cell(0, 10, f"Q{i+1}: {q['question']}", 0, 1)
        for choice in q['choices']:
            pdf.cell(0, 10, f" - {choice}", 0, 1)
        if include_answers:
            pdf.cell(0, 10, f"Correct Answer: {q['correct_answer']}", 0, 1)
            pdf.cell(0, 10, f"Explanation: {q['explanation']}", 0, 1)
        pdf.ln(10)

    return pdf.output(dest="S").encode("latin1")

def generate_docx(questions, include_answers=True):
    """Generates a DOCX file with exam questions."""
    doc = Document()
    doc.add_heading('Generated Exam', level=1)
    for i, q in enumerate(questions):
        doc.add_heading(f"Q{i+1}: {q['question']}", level=2)
        for choice in q['choices']:
            doc.add_paragraph(choice, style='List Bullet')
        if include_answers:
            doc.add_paragraph(f"Correct Answer: {q['correct_answer']}")
            doc.add_paragraph(f"Explanation: {q['explanation']}")
        doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --------------------------- Main Application ---------------------------

def main():
    st.title("📝 Exam Creator")
    st.markdown(f"**Version:** {__version__}")

    if "client" not in st.session_state:
        st.session_state.client = None

    if "generated_questions" not in st.session_state:
        st.session_state.generated_questions = []

    st.sidebar.title("Configuration")
    api_key = st.sidebar.text_input("OpenAI API Key", type="password")
    model = st.sidebar.selectbox("Model", ["gpt-4", "gpt-3.5-turbo"], index=0)

    if api_key:
        try:
            st.session_state.client = OpenAI(api_key=api_key)
            st.sidebar.success("API Key configured successfully!")
        except Exception as e:
            st.sidebar.error(f"Error initializing OpenAI client: {e}")

    mode = st.sidebar.radio("Choose Mode", ["Upload & Generate", "Take Quiz", "Download Exam"])

    if mode == "Upload & Generate":
        st.subheader("Upload File to Generate Exam")
        uploaded_file = st.file_uploader("Upload PDF or DOCX", type=["pdf", "docx"])

        if uploaded_file:
            if uploaded_file.type == "application/pdf":
                content = extract_text_from_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                content = extract_text_from_docx(uploaded_file)
            else:
                st.error("Unsupported file type.")
                return

            if content:
                st.text_area("Extracted Content", content[:500] + "...", height=200)
                if st.button("Generate Questions"):
                    client = st.session_state.client
                    chunks = chunk_text(content)
                    questions = []
                    for chunk in chunks:
                        response, error = generate_mc_questions(client, chunk, model)
                        if error:
                            st.error(error)
                            break
                        parsed_questions, parse_error = parse_generated_questions(response)
                        if parse_error:
                            st.error(parse_error)
                            break
                        questions.extend(parsed_questions)
                        if len(questions) >= 20:
                            break
                    st.session_state.generated_questions = questions[:20]
                    st.success(f"Generated {len(questions[:20])} questions!")

    elif mode == "Take Quiz":
        questions = st.session_state.generated_questions
        if questions:
            for i, q in enumerate(questions):
                st.write(f"Q{i+1}: {q['question']}")
                st.radio("Choose an answer:", q['choices'], key=f"q_{i}")
        else:
            st.warning("No questions generated yet.")

    elif mode == "Download Exam":
        questions = st.session_state.generated_questions
        if questions:
            format_option = st.radio("Choose Format", ["PDF", "DOCX"])
            include_answers = st.checkbox("Include Answers", value=True)
            if st.button("Download"):
                if format_option == "PDF":
                    file_data = generate_pdf(questions, include_answers)
                    file_name = "exam.pdf"
                else:
                    file_data = generate_docx(questions, include_answers)
                    file_name = "exam.docx"
                st.download_button("Download", data=file_data, file_name=file_name)
        else:
            st.warning("No questions generated yet.")

if __name__ == "__main__":
    main()
